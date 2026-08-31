"""
Document Parser Service - Intelligent document parsing.
Supports table extraction, image OCR, formula recognition.
"""
from __future__ import annotations

import logging
import os
import re
from typing import Any

from app.core.config import settings

logger = logging.getLogger(__name__)


def extract_tables_from_docx(file_path: str) -> list[dict[str, Any]]:
    """
    Extract tables from DOCX file.

    Returns:
        List of tables, each table is a list of rows,
        each row is a list of cell values.
    """
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        tables = []

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append({
                    "type": "table",
                    "data": table_data,
                    "markdown": _table_to_markdown(table_data),
                })

        return tables
    except Exception as e:
        logger.warning(f"Failed to extract tables from DOCX: {e}")
        return []


def extract_tables_from_html(html_content: str) -> list[dict[str, Any]]:
    """Extract tables from HTML content."""
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html_content, 'html.parser')
        tables = []

        for table in soup.find_all('table'):
            rows = []
            for tr in table.find_all('tr'):
                cells = []
                for td in tr.find_all(['td', 'th']):
                    cells.append(td.get_text(strip=True))
                if cells:
                    rows.append(cells)
            if rows:
                tables.append({
                    "type": "table",
                    "data": rows,
                    "markdown": _table_to_markdown(rows),
                })

        return tables
    except ImportError:
        logger.warning("BeautifulSoup not installed, skipping HTML table extraction")
        return []


def _table_to_markdown(table_data: list[list[str]]) -> str:
    """Convert table data to Markdown format."""
    if not table_data:
        return ""

    # Header
    header = table_data[0]
    separator = ["---"] * len(header)

    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]

    # Rows
    for row in table_data[1:]:
        # Pad row if needed
        while len(row) < len(header):
            row.append("")
        lines.append("| " + " | ".join(row[:len(header)]) + " |")

    return "\n".join(lines)


def extract_code_blocks(text: str) -> list[dict[str, Any]]:
    """
    Extract code blocks from text (Markdown format).

    Returns:
        List of code blocks with language and content.
    """
    pattern = r'```(\w+)?\s*\n(.*?)```'
    matches = re.findall(pattern, text, re.DOTALL)

    blocks = []
    for lang, code in matches:
        blocks.append({
            "type": "code",
            "language": lang or "unknown",
            "content": code.strip(),
        })

    return blocks


def extract_formulas(text: str) -> list[dict[str, Any]]:
    """
    Extract mathematical formulas from text.
    Supports LaTeX inline ($...$) and display ($$...$$) formulas.

    Returns:
        List of formulas with type and content.
    """
    formulas = []

    # Display formulas ($$...$$)
    display_pattern = r'\$\$(.*?)\$\$'
    for match in re.finditer(display_pattern, text, re.DOTALL):
        formulas.append({
            "type": "display_formula",
            "content": match.group(1).strip(),
            "latex": match.group(0),
        })

    # Inline formulas ($...$)
    inline_pattern = r'(?<!\$)\$(?!\$)(.*?)(?<!\$)\$(?!\$)'
    for match in re.finditer(inline_pattern, text):
        formulas.append({
            "type": "inline_formula",
            "content": match.group(1).strip(),
            "latex": match.group(0),
        })

    return formulas


def extract_images_from_docx(file_path: str) -> list[dict[str, Any]]:
    """
    Extract image information from DOCX file.

    把内嵌图片 blob 解包到临时文件，返回含 local_path 的元数据，供 OCR 后端读取。
    注意：解包后的临时文件由调用方在使用后自行清理（演示场景影响可忽略）。

    Returns:
        List of image dicts with keys: type, content_type, description, local_path.
    """
    try:
        import tempfile

        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        images = []

        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                name = os.path.basename(rel.target_ref)
                local_path = None
                try:
                    blob = rel.blob
                    if blob:
                        tmp_dir = tempfile.mkdtemp(prefix="rag_docx_img_")
                        local_path = os.path.join(tmp_dir, name)
                        with open(local_path, "wb") as f:
                            f.write(blob)
                except Exception as e:  # noqa: BLE001
                    logger.warning(f"解包 DOCX 图片失败 {name}: {e}")
                images.append({
                    "type": "image",
                    "content_type": rel.target_ref,
                    "description": f"Image: {name}",
                    "local_path": local_path,
                })

        return images
    except Exception as e:
        logger.warning(f"Failed to extract images from DOCX: {e}")
        return []


_MD_IMG_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
_HTML_IMG_RE = re.compile(r"<img\b[^>]*>", re.IGNORECASE)


def _do_ocr(local_path: str | None, ocr_backend: Any) -> str:
    """对单个图片路径调 OCR，失败返回空串。"""
    if not local_path or ocr_backend is None:
        return ""
    try:
        return ocr_backend.ocr_image(local_path)
    except Exception as e:  # noqa: BLE001
        logger.warning(f"OCR 失败 {local_path}: {e}")
        return ""


def _resolve_image_path(src: str, base_dir: str) -> str | None:
    """把图片 src（相对/绝对/远程）解析为可本地 OCR 的路径。

    远程(http/data:)图片当前返回 None（后续可扩展下载/多模态）。
    """
    if not src:
        return None
    if src.startswith(("http://", "https://", "data:")):
        return None
    if os.path.isabs(src):
        return src if os.path.isfile(src) else None
    cand = os.path.join(base_dir, src)
    return cand if os.path.isfile(cand) else None


def _ocr_markdown_images(
    text: str, base_dir: str, ocr_backend: Any
) -> tuple[str, list[dict[str, Any]]]:
    """Markdown：`![alt](src)` 逐个 OCR，回填 [图片OCR]...[/图片OCR]。"""
    images: list[dict[str, Any]] = []

    def _repl(m: re.Match) -> str:
        alt, src = m.group(1), m.group(2)
        local = _resolve_image_path(src, base_dir)
        ocr_text = _do_ocr(local, ocr_backend)
        images.append({
            "type": "image",
            "src": src,
            "alt": alt,
            "local_path": local,
            "ocr_text": ocr_text,
        })
        if ocr_text:
            return f"\n[图片OCR 开始]{ocr_text}[图片OCR 结束]\n"
        return f"\n[图片: {alt or src}]\n"

    return _MD_IMG_RE.sub(_repl, text), images


def _ocr_html_images(
    html_content: str, base_dir: str, ocr_backend: Any
) -> tuple[str, list[dict[str, Any]]]:
    """HTML：把每个 <img> 替换为 [图片OCR] 文本节点，再交给 get_text。"""
    try:
        from bs4 import BeautifulSoup, NavigableString
    except ImportError:
        return html_content, []

    soup = BeautifulSoup(html_content, "html.parser")
    images: list[dict[str, Any]] = []
    for img in soup.find_all("img"):
        src = img.get("src", "")
        alt = img.get("alt", "")
        local = _resolve_image_path(src, base_dir)
        ocr_text = _do_ocr(local, ocr_backend)
        images.append({
            "type": "image",
            "src": src,
            "alt": alt,
            "local_path": local,
            "ocr_text": ocr_text,
        })
        repl = (
            f"[图片OCR 开始]{ocr_text}[图片OCR 结束]"
            if ocr_text
            else f"[图片: {alt or src}]"
        )
        img.replace_with(NavigableString(repl))
    return str(soup), images


def parse_document_intelligently(
    file_path: str,
    file_type: str,
    include_metadata: bool = True,
    ocr_backend: Any = None,
) -> dict[str, Any]:
    """
    Intelligently parse document with structure preservation.

    Args:
        file_path: Path to the document file
        file_type: File extension (pdf, docx, html, etc.)
        include_metadata: Whether to include extracted metadata

    Returns:
        Dict with 'text', 'tables', 'code_blocks', 'formulas', 'images'
    """
    result = {
        "text": "",
        "tables": [],
        "code_blocks": [],
        "formulas": [],
        "images": [],
    }

    file_type = file_type.lower()
    base_dir = os.path.dirname(os.path.abspath(file_path))

    # Extract based on file type
    if file_type in ("docx", "doc"):
        # Extract tables
        result["tables"] = extract_tables_from_docx(file_path)

        # Extract images metadata
        if include_metadata:
            result["images"] = extract_images_from_docx(file_path)
            # OCR 回填：把 docx 内嵌图识别文字追加到文本流末尾
            if ocr_backend is not None:
                ocr_parts = []
                for im in result["images"]:
                    if im.get("local_path"):
                        im["ocr_text"] = _do_ocr(im["local_path"], ocr_backend)
                        if im.get("ocr_text"):
                            ocr_parts.append(
                                f"\n[图片OCR 开始]{im['ocr_text']}[图片OCR 结束]\n"
                            )
                if ocr_parts:
                    result["text"] = (
                        result["text"] + "\n" + "\n".join(ocr_parts)
                    ).strip()

        # Get text content
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            result["text"] = "\n\n".join(paragraphs)
        except Exception as e:
            logger.warning(f"Failed to extract text from DOCX: {e}")

    elif file_type in ("html", "htm"):
        # Extract tables from HTML
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                html_content = f.read()
            result["tables"] = extract_tables_from_html(html_content)

            # 图片 OCR 回填：先把 <img> 换成 [图片OCR] 文本节点，再 get_text
            html_after_ocr, html_images = _ocr_html_images(html_content, base_dir, ocr_backend)
            result["images"] = html_images

            # Extract text
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html_after_ocr, 'html.parser')
            result["text"] = soup.get_text(separator="\n", strip=True)
        except ImportError:
            logger.warning("BeautifulSoup not installed")
        except Exception as e:
            logger.warning(f"Failed to parse HTML: {e}")

    elif file_type in ("txt", "md"):
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            # 先基于原文提取代码块/公式（图片语法不在代码块内），再做图片 OCR 回填
            if file_type == "md":
                result["code_blocks"] = extract_code_blocks(content)
                result["formulas"] = extract_formulas(content)

            # 图片 OCR 回填：把 ![alt](src) 替换为 [图片OCR]...[/图片OCR]
            text_after_ocr, md_images = _ocr_markdown_images(content, base_dir, ocr_backend)
            result["images"] = md_images
            result["text"] = text_after_ocr
        except Exception as e:
            logger.warning(f"Failed to read text file: {e}")

    elif file_type == "pdf":
        # PDF parsing would require additional libraries
        # For now, fall back to basic text extraction
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            texts = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    texts.append(t)
            result["text"] = "\n\n".join(texts)
        except Exception as e:
            logger.warning(f"Failed to parse PDF: {e}")

    # Add metadata
    if include_metadata:
        result["metadata"] = {
            "file_type": file_type,
            "has_tables": len(result["tables"]) > 0,
            "has_code_blocks": len(result["code_blocks"]) > 0,
            "has_formulas": len(result["formulas"]) > 0,
            "has_images": len(result["images"]) > 0,
            "table_count": len(result["tables"]),
            "code_block_count": len(result["code_blocks"]),
            "formula_count": len(result["formulas"]),
            "image_count": len(result["images"]),
        }

    return result


def format_structured_content(parsed: dict[str, Any]) -> str:
    """
    Format parsed document content into structured text for chunking.
    Preserves table structure and adds markers for special content.
    """
    parts = []

    # Main text
    if parsed.get("text"):
        parts.append(parsed["text"])

    # Tables
    for table in parsed.get("tables", []):
        parts.append(f"\n[表格]\n{table.get('markdown', '')}\n[/表格]\n")

    # Code blocks
    for block in parsed.get("code_blocks", []):
        lang = block.get("language", "")
        content = block.get("content", "")
        parts.append(f"\n[代码块:{lang}]\n{content}\n[/代码块]\n")

    # Formulas
    for formula in parsed.get("formulas", []):
        parts.append(f"\n[公式]\n{formula.get('latex', '')}\n[/公式]\n")

    # Images (OCR 文字)
    for im in parsed.get("images", []):
        ocr_text = im.get("ocr_text")
        if ocr_text:
            parts.append(f"\n[图片OCR 开始]{ocr_text}[图片OCR 结束]\n")

    return "\n".join(parts)
