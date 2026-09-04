"""
Documents API: file upload, status tracking, deletion.
"""
import hashlib
import logging
import os
import tempfile
import uuid
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException, Query, Request, UploadFile, File, status
from typing import Annotated

logger = logging.getLogger(__name__)
from sqlalchemy import select, func, and_
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import joinedload

from app.core.config import settings
from app.core.database import get_db
from app.core.dependencies import get_current_user
from app.core.security import md5_hash
from app.models.audit_log import AuditLog
from app.models.document import Document, DocumentStatus
from app.models.user import User
from app.schemas.common import DocumentInfoResponse, DocumentListResponse

router = APIRouter()

# Import limiter from core
from app.core.limiter import limiter

# 流式写盘时的单次读取大小
UPLOAD_CHUNK_SIZE = 1024 * 1024  # 1 MiB


@router.post("/upload", status_code=status.HTTP_201_CREATED)
@limiter.limit("20/minute")
async def upload_document(
    request: Request,
    kb_id: int,
    upload_file: Annotated[UploadFile, File(...)],
    chunk_strategy: str = "fixed_token",
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    接口说明：上传文档到指定知识库，校验格式与大小后落盘并建库记录，随后触发异步处理流水线（解析→切片→向量化→写入 Milvus），并写入审计日志。
    方法路径：POST /api/documents/upload
    鉴权要求：已登录且拥有该 KB 的 upload 及以上权限
    表单参数：kb_id(int,必填), upload_file(File,必填), chunk_strategy(str,选填,默认 fixed_token)
    响应字段：DocumentInfoResponse{id,filename,original_filename,file_size,file_type,status,chunk_strategy,...}
    错误码：400 文件类型不允许/大小超限; 401 未登录; 403 无上传权限; 429 上传频率超限(20/min)
    """
    import traceback
    try:
        from app.api.knowledge_bases import _check_kb_permission

        ip = request.client.host if request.client else ""
        ua = request.headers.get("user-agent", "")

        # Permission check (need 'upload' level)
        kb, _ = await _check_kb_permission(db, kb_id, current_user, min_level="upload")

        # Validate extension
        original_filename = upload_file.filename or "unknown"
        ext = original_filename.rsplit(".", 1)[-1].lower() if "." in original_filename else ""
        if ext not in settings.upload_allowed_extensions:
            raise HTTPException(
                400,
                f"File type '{ext}' not allowed. Allowed: {settings.upload_allowed_extensions}",
            )

        # Stream to disk while enforcing the size limit.
        # 原实现 await upload_file.read() 把整个文件读进内存：上传上限
        # 100MB 时，10 个并发上传就是 1GB 常驻内存，且超限文件要完整
        # 收完才被发现。改为边收边写，超限立即中断，内存占用恒定为
        # 一个块大小。文件名依赖内容哈希，先写临时文件再改名。
        upload_dir = settings.UPLOAD_DIR
        # 临时文件写进 UPLOAD_DIR 下的 .staging 子目录，不落在正式文件堆里：
        # 进程被 kill 时 finally 不执行，混放的 .upload 残骸只能靠扩展名在
        # 正式文件里翻找，而独立子目录的清理作用域是确定的（见 main.py 的
        # 启动清理）。子目录与正式目录同一文件系统，os.replace 仍是原子改名。
        staging_dir = settings.upload_staging_dir
        os.makedirs(upload_dir, exist_ok=True)
        os.makedirs(staging_dir, exist_ok=True)

        hasher = hashlib.md5()
        file_size = 0
        tmp_fd, tmp_path = tempfile.mkstemp(dir=staging_dir, suffix=".upload")
        save_path = None
        try:
            with os.fdopen(tmp_fd, "wb") as out:
                while True:
                    chunk = await upload_file.read(UPLOAD_CHUNK_SIZE)
                    if not chunk:
                        break
                    file_size += len(chunk)
                    if file_size > settings.upload_max_bytes:
                        raise HTTPException(
                            413,
                            f"File too large. Max: {settings.UPLOAD_MAX_SIZE_MB}MB",
                        )
                    hasher.update(chunk)
                    out.write(chunk)

            # Generate unique filename (content hash + random salt)
            file_hash = md5_hash(hasher.hexdigest() + str(uuid.uuid4()))
            stored_filename = f"{file_hash}.{ext}"
            save_path = os.path.join(upload_dir, stored_filename)
            os.replace(tmp_path, save_path)
            tmp_path = None  # 已改名，不要在 finally 里删掉正式文件
        finally:
            if tmp_path is not None and os.path.exists(tmp_path):
                try:
                    os.unlink(tmp_path)
                except OSError as e:
                    logger.warning(f"[upload] Failed to remove temp file {tmp_path}: {e}")

        # Create DB record
        doc = Document(
            kb_id=kb_id,
            filename=stored_filename,
            original_filename=original_filename,
            file_size=file_size,
            file_type=ext,
            uploader_id=current_user.id,
            chunk_strategy=chunk_strategy,
            chunk_count=0,
            chunk_params={
                "chunk_size": settings.DEFAULT_CHUNK_SIZE,
                "overlap": settings.DEFAULT_CHUNK_OVERLAP,
            },
            status=DocumentStatus.pending,
        )
        db.add(doc)
        await db.commit()
        await db.refresh(doc)

        # Audit log
        audit = AuditLog(
            user_id=current_user.id,
            action="upload",
            resource_type="document",
            resource_id=doc.id,
            detail={
                "kb_id": kb_id,
                "original_filename": original_filename,
                "file_size": file_size,
                "chunk_strategy": chunk_strategy,
            },
            ip_address=ip,
            user_agent=ua,
        )
        db.add(audit)
        await db.commit()

        resp = DocumentInfoResponse.model_validate(doc)
        resp.uploader_name = current_user.real_name

        # Trigger background processing pipeline (non-blocking)
        _trigger_processing_pipeline(doc.id, current_user)

        return resp

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[upload] Unexpected error: {e}\n{traceback.format_exc()}")
        # Don't expose internal error details to client
        raise HTTPException(status_code=500, detail="文件上传失败，请稍后重试")


_background_tasks: set = set()
_task_counter = 0


def _trigger_processing_pipeline(document_id: int, user):
    """Fire-and-forget: parse -> chunk -> embed -> store to Milvus."""
    import asyncio
    global _task_counter
    _task_counter += 1
    task_id = _task_counter

    async def _run_pipeline():
        logger.info(f"[pipeline] Task {task_id} started for document {document_id}")
        try:
            from app.services.kb_service import process_document_async
            await process_document_async(document_id, user.id)
            logger.info(f"[pipeline] Task {task_id} completed for document {document_id}")
        except Exception as e:
            import traceback
            logger.error(f"[pipeline] Task {task_id} FAILED for document {document_id}: {e}\n{traceback.format_exc()}")
            # Update document status to failed in DB
            try:
                from app.core.database import AsyncSessionLocal
                from sqlalchemy import text
                async with AsyncSessionLocal() as db:
                    await db.execute(
                        text("UPDATE documents SET status='failed', error_msg=:err WHERE id=:did"),
                        {"err": f"处理失败: {str(e)[:400]}", "did": document_id},
                    )
                    await db.commit()
            except Exception as db_err:
                logger.error(f"[pipeline] Failed to update doc {document_id} error status: {db_err}")

    task = asyncio.create_task(_run_pipeline())
    _background_tasks.add(task)
    task.add_done_callback(lambda t: _background_tasks.discard(t))
    logger.info(f"[pipeline] Task {task_id} queued for document {document_id}")


@router.get("", response_model=DocumentListResponse)  # 文档列表
async def list_documents(
    kb_id: int | None = None,
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    status_filter: str | None = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    接口说明：分页查询文档列表，支持按知识库与处理状态过滤，并附带上传者姓名。
    方法路径：GET /api/documents
    鉴权要求：已登录任意角色用户（指定 kb_id 时自动校验该 KB 访问权限）
    请求参数：kb_id(int,选填), page(int,默认1), page_size(int 1-100,默认20), status_filter(str,选填)
    响应字段：DocumentListResponse{total, items[DocumentInfoResponse]}
    错误码：401 未登录; 403 无 KB 访问权限
    """
    filters = [Document.is_deleted == False]

    if kb_id is not None:
        from app.api.knowledge_bases import _check_kb_permission
        await _check_kb_permission(db, kb_id, current_user)
        filters.append(Document.kb_id == kb_id)

    if status_filter:
        filters.append(Document.status == status_filter)

    count_q = select(func.count(Document.id)).where(*filters)
    total_result = await db.execute(count_q)
    total = total_result.scalar() or 0

    offset = (page - 1) * page_size
    result = await db.execute(
        select(Document)
        .options(joinedload(Document.uploader))
        .where(*filters)
        .order_by(Document.created_at.desc())
        .offset(offset)
        .limit(page_size)
    )
    docs = result.scalars().unique().all()

    items = []
    for d in docs:
        resp = DocumentInfoResponse.model_validate(d)
        resp.uploader_name = d.uploader.real_name if d.uploader else ""
        items.append(resp)

    return DocumentListResponse(total=total, items=items)


@router.get("/{document_id}", response_model=DocumentInfoResponse)  # 文档详情
async def get_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    接口说明：获取单个文档的详细信息（含上传者姓名）。
    方法路径：GET /api/documents/{document_id}
    鉴权要求：已登录任意角色用户（文档本身不做 KB 级权限强校验，仅校验存在性）
    路径参数：document_id(int,必填)
    响应字段：DocumentInfoResponse{id,original_filename,file_size,file_type,status,chunk_count,uploader_name,...}
    错误码：401 未登录; 404 文档不存在/已删除
    """
    # Use joinedload to avoid N+1 query
    result = await db.execute(
        select(Document)
        .options(joinedload(Document.uploader))
        .where(Document.id == document_id)
    )
    doc = result.scalars().unique().first()
    if doc is None or doc.is_deleted:
        raise HTTPException(status_code=404, detail="Document not found")
    resp = DocumentInfoResponse.model_validate(doc)
    resp.uploader_name = doc.uploader.real_name if doc.uploader else ""
    return resp


@router.delete("/{document_id}")  # 删除文档：软删
async def delete_document(
    document_id: int,
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    接口说明：软删除文档（置 is_deleted=True），级联软删除其切片，并异步清理 Milvus 中对应向量，写入审计日志。
    方法路径：DELETE /api/documents/{document_id}
    鉴权要求：已登录且拥有该文档所属 KB 的 upload 及以上权限
    路径参数：document_id(int,必填)
    响应字段：detail("Document soft-deleted"), id
    错误码：401 未登录; 403 无上传权限; 404 文档不存在/已删除
    """
    ip = request.client.host if request.client else ""
    ua = request.headers.get("user-agent", "")

    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if doc is None or doc.is_deleted:
        raise HTTPException(404, "Document not found")

    # Permission check
    from app.api.knowledge_bases import _check_kb_permission
    await _check_kb_permission(db, doc.kb_id, current_user, min_level="upload")

    doc.is_deleted = True
    doc.deleted_at = func.now()

    audit = AuditLog(
        user_id=current_user.id,
        action="delete",
        resource_type="document",
        resource_id=document_id,
        detail={"filename": doc.original_filename, "kb_id": doc.kb_id},
        ip_address=ip,
        user_agent=ua,
    )
    db.add(audit)

    # Also soft-delete associated chunks and clean Milvus vectors
    from app.models.chunk import Chunk
    from app.services.milvus_service import delete_by_chunk_ids

    # Get chunk IDs for Milvus cleanup
    chunk_ids_result = await db.execute(
        select(Chunk.id).where(Chunk.document_id == document_id, Chunk.is_deleted == False)
    )
    chunk_ids_to_delete = [row[0] for row in chunk_ids_result.fetchall()]

    await db.execute(
        Chunk.__table__.update().where(Chunk.document_id == document_id).values(is_deleted=True)
    )

    await db.commit()

    # Clean up Milvus vectors (fire-and-forget, don't block response)
    if chunk_ids_to_delete:
        try:
            await delete_by_chunk_ids(doc.kb_id, chunk_ids_to_delete)
        except Exception as e:
            import logging
            logging.getLogger(__name__).warning(f"[delete] Milvus cleanup failed for doc {document_id}: {e}")

    # 知识库内容已删除（片段从向量库移除，检索结果会变）→ 使该 KB 答案缓存失效。
    try:
        from app.services.answer_cache import bump_kb_epoch

        await bump_kb_epoch(doc.kb_id)
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning(f"[delete] bump_kb_epoch failed (kb={doc.kb_id}): {e}")

    return {"detail": "Document soft-deleted", "id": doc.id}


@router.get("/{document_id}/preview")  # 文档在线预览：多格式
async def preview_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    接口说明：在线预览文档内容：PDF 返回 base64 供前端渲染，DOCX 转 HTML，TXT/MD 返回纯文本，CSV/XLSX 转 HTML 表格（预览限前 100 行）。
    方法路径：GET /api/documents/{document_id}/preview
    鉴权要求：已登录且拥有该文档所属 KB 的 read 及以上权限
    路径参数：document_id(int,必填)
    响应字段：{type("pdf"/"html"/"text"), data, filename}
    错误码：401 未登录; 403 无读取权限; 404 文档/文件不存在; 400 不支持的格式; 500 预览解析失败
    """
    from fastapi.responses import Response
    import base64
    import os

    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if doc is None or doc.is_deleted:
        raise HTTPException(status_code=404, detail="Document not found")

    # Permission check
    from app.api.knowledge_bases import _check_kb_permission
    await _check_kb_permission(db, doc.kb_id, current_user)

    file_path = os.path.join(settings.UPLOAD_DIR, doc.filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found on disk")

    file_type = doc.file_type.lower()

    # PDF: return as base64 for frontend rendering
    if file_type == "pdf":
        with open(file_path, "rb") as f:
            content = f.read()
        return {
            "type": "pdf",
            "data": base64.b64encode(content).decode("utf-8"),
            "filename": doc.original_filename,
        }

    # DOCX: convert to HTML for preview
    elif file_type in ("docx", "doc"):
        try:
            from docx import Document as DocxDocument
            docx_doc = DocxDocument(file_path)
            html_parts = []
            for para in docx_doc.paragraphs:
                if para.text.strip():
                    style = para.style.name if para.style else ""
                    if style.startswith("Heading"):
                        level = style.replace("Heading ", "")
                        html_parts.append(f"<h{level}>{para.text}</h{level}>")
                    else:
                        html_parts.append(f"<p>{para.text}</p>")
            return {
                "type": "html",
                "data": "\n".join(html_parts),
                "filename": doc.original_filename,
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to preview DOCX: {str(e)}")

    # TXT/MD: return as text
    elif file_type in ("txt", "md"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return {
            "type": "text",
            "data": content,
            "filename": doc.original_filename,
        }

    # CSV: convert to HTML table
    elif file_type == "csv":
        try:
            import csv
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                reader = csv.reader(f)
                rows = list(reader)

            if not rows:
                return {"type": "text", "data": "空文件", "filename": doc.original_filename}

            # Build HTML table
            html_parts = ["<table border='1' cellpadding='8' cellspacing='0' style='border-collapse:collapse;width:100%'>"]
            # Header
            html_parts.append("<tr>")
            for cell in rows[0]:
                html_parts.append(f"<th style='background:#f5f5f5;text-align:left'>{cell}</th>")
            html_parts.append("</tr>")
            # Rows
            for row in rows[1:100]:  # Limit to 100 rows for preview
                html_parts.append("<tr>")
                for cell in row:
                    html_parts.append(f"<td>{cell}</td>")
                html_parts.append("</tr>")
            html_parts.append("</table>")

            if len(rows) > 101:
                html_parts.append(f"<p style='color:#999;margin-top:8px'>仅显示前100行，共{len(rows)-1}行数据</p>")

            return {
                "type": "html",
                "data": "\n".join(html_parts),
                "filename": doc.original_filename,
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to preview CSV: {str(e)}")

    # XLSX/XLS: convert to HTML table
    elif file_type in ("xlsx", "xls"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            ws = wb.active

            html_parts = ["<table border='1' cellpadding='8' cellspacing='0' style='border-collapse:collapse;width:100%'>"]
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                if row_count == 0:
                    # Header
                    html_parts.append("<tr>")
                    for cell in row:
                        html_parts.append(f"<th style='background:#f5f5f5;text-align:left'>{cell or ''}</th>")
                    html_parts.append("</tr>")
                elif row_count < 100:  # Limit to 100 rows
                    html_parts.append("<tr>")
                    for cell in row:
                        html_parts.append(f"<td>{cell if cell is not None else ''}</td>")
                    html_parts.append("</tr>")
                row_count += 1
            html_parts.append("</table>")

            if row_count > 100:
                html_parts.append(f"<p style='color:#999;margin-top:8px'>仅显示前100行，共{row_count}行数据</p>")

            wb.close()
            return {
                "type": "html",
                "data": "\n".join(html_parts),
                "filename": doc.original_filename,
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to preview Excel: {str(e)}")

    # Unsupported format
    else:
        raise HTTPException(status_code=400, detail=f"Preview not supported for .{file_type} files")
